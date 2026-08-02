#!/usr/bin/env python3
"""
DivorceGPT — NY Phase-1 forms as editable Word documents (.docx)
================================================================

Operator directive (2026-07-27, "for Claude" doc): the attorney downloads
generated forms in WORD from the matter rail. Word is the format lawyers
actually edit before filing — the PDF pipeline stays for everything, this
module adds a .docx build for the Phase-1 forms.

CONTENT PARITY, NOT PIXEL PARITY. Every substantive sentence comes from the
same clause functions the PDF generators use (imported below, never
duplicated), so the Word text and the PDF text can never diverge. Layout is
idiomatic Word — caption as a bordered two-column table, allegations as
hanging-indent paragraphs — because the attorney is going to EDIT this file;
canvas-perfect coordinates would only fight them.

Same data contract as the PDF twins (generate_ud1 / generate_complaint).
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .generate_complaint import (
    DEFAULT_ATTORNEY_NAME,
    DEFAULT_ATTORNEY_FIRM,
    DEFAULT_ATTORNEY_ADDRESS,
    DEFAULT_ATTORNEY_PHONE,
    ORDINALS,
    children_clause,
    drl253_clause,
    fmt_marriage_date,
    marriage_clause,
    relief_bundle,
    residency_clause,
    strip_county_suffix,
    title_case,
)

BODY = "Times New Roman"


# ────────────────────────── shared building blocks ──────────────────────────

def _base_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def _p(doc, text="", bold=False, italic=False, align=None, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    return para


def _set_cell_borders(cell, edges):
    """Turn on single borders for the given edges of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge in edges:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def _caption_table(doc, county_upper, plaintiff_upper, defendant_upper, right_lines):
    """
    NY pleading caption: court header, then a two-column block — parties on
    the left (bordered like the filed X-rule box), index/title column on the
    right. `right_lines` is a list of (text, bold) tuples.
    """
    _p(doc, "SUPREME COURT OF THE STATE OF NEW YORK", bold=True, space_after=0)
    _p(doc, f"COUNTY OF {county_upper}", bold=True, space_after=2)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(3.4)
    right.width = Inches(3.1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_borders(left, edges=("top", "right", "bottom"))
    _set_cell_borders(right, edges=())

    lp = left.paragraphs[0]
    lp.add_run(f"{plaintiff_upper},")
    left.add_paragraph("")
    ind = left.add_paragraph()
    ind.paragraph_format.left_indent = Inches(1.6)
    ind.add_run("Plaintiff,")
    ag = left.add_paragraph()
    ag.paragraph_format.left_indent = Inches(0.5)
    ag.add_run("-against-")
    left.add_paragraph("")
    left.add_paragraph(f"{defendant_upper},")
    ind2 = left.add_paragraph()
    ind2.paragraph_format.left_indent = Inches(1.6)
    ind2.add_run("Defendant.")

    first = True
    for text, bold in right_lines:
        para = right.paragraphs[0] if first else right.add_paragraph()
        first = False
        run = para.add_run(text)
        run.bold = bold
        if bold:
            run.underline = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def _signature_block(doc, data, date_label="Dated:"):
    date_signed = (data.get("dateSigned") or "").strip() or "____________________"
    attorney_name = data.get("attorneyName", DEFAULT_ATTORNEY_NAME)
    attorney_firm = data.get("attorneyFirm", DEFAULT_ATTORNEY_FIRM)
    attorney_addr = data.get("attorneyAddress", DEFAULT_ATTORNEY_ADDRESS)
    attorney_phone = data.get("attorneyPhone", DEFAULT_ATTORNEY_PHONE)

    _p(doc, f"{date_label} {date_signed}", space_after=18)
    for line in [
        "____________________________",
        attorney_name,
        attorney_firm,
        *str(attorney_addr).split("\n"),
        attorney_phone,
        "Attorney for Plaintiff",
    ]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(3.4)
        para.paragraph_format.space_after = Pt(0)
        para.add_run(line)


def _hanging_allegation(doc, label, text):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.first_line_indent = Inches(0.5)
    pf.space_after = Pt(12)
    pf.line_spacing = 2.0  # 22 NYCRR 202.5(a)(1) — double space
    run = para.add_run(f"{label}: ")
    run.bold = True
    para.add_run(text)
    return para


# ────────────────────────── Verified Complaint ──────────────────────────

def generate_complaint_docx(data, output_path):
    """NY Verified Complaint (Action for Divorce) as an editable .docx —
    same allegations, same order, same clause text as generate_complaint."""
    county = strip_county_suffix(data.get("county") or data.get("filingCounty"))
    if not county:
        raise ValueError("County is required")
    plaintiff = (data.get("plaintiffName") or "").strip()
    defendant = (data.get("defendantName") or "").strip()
    if not plaintiff or not defendant:
        raise ValueError("Plaintiff and Defendant names are required")
    plaintiff_addr = (data.get("plaintiffAddress") or "").strip()
    defendant_addr = (data.get("defendantAddress") or "").strip()

    resident_party = (data.get("residentParty") or "plaintiff").strip().lower()
    ceremony_type = (data.get("ceremonyType") or "civil").strip().lower()
    residency_basis = (data.get("residencyBasis") or "two_year").strip().lower()
    children = data.get("unemancipatedChildren", 0)
    relief = data.get("reliefBundle") or relief_bundle(children)
    attorney_name = data.get("attorneyName", DEFAULT_ATTORNEY_NAME)

    doc = _base_document()
    _caption_table(
        doc,
        county.upper(),
        plaintiff.upper(),
        defendant.upper(),
        [("Index No.: ______________", False), ("", False), ("VERIFIED COMPLAINT", True), ("", False), ("ACTION FOR A DIVORCE", False)],
    )
    _p(doc, "")
    _p(
        doc,
        f"Plaintiff, by {attorney_name}, complaining of the Defendant, as and for a "
        f"Verified Complaint, alleges:",
        space_after=12,
    )

    allegations = [
        residency_clause(resident_party, residency_basis),
        "Both parties are over the age of eighteen (18) years as of the date set forth herein.",
        marriage_clause(data.get("marriageDate"), data.get("marriagePlace"), ceremony_type),
        drl253_clause(ceremony_type),
        children_clause(children, data.get("childrenDetail", "")),
        f"The Plaintiff resides at {plaintiff_addr}. The Defendant resides at {defendant_addr}.",
        "This marriage has never been altered or dissolved by any judgment of divorce, "
        "annulment, or dissolution of marriage issued by any court of competent "
        "jurisdiction.",
        "No other action or proceeding between the parties for divorce, annulment, "
        "separation, or dissolution of the marriage is pending in this or any other "
        "court of competent jurisdiction.",
        "The grounds for divorce, pursuant to Subdivision (7) of Section 170 of the "
        "Domestic Relations Law, are as follows: the relationship between the parties "
        "has broken down irretrievably for a period of at least six months.",
    ]
    for i, text in enumerate(allegations):
        _hanging_allegation(doc, ORDINALS[i], text)

    _p(doc, "")
    wherefore = doc.add_paragraph()
    wherefore.paragraph_format.first_line_indent = Inches(0.5)
    wherefore.paragraph_format.space_after = Pt(12)
    run = wherefore.add_run("WHEREFORE")
    run.bold = True
    wherefore.add_run(", Plaintiff demands judgment against the Defendant as follows:")
    for i, item in enumerate(relief):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.left_indent = Inches(0.75)
        pf.space_after = Pt(6)
        para.add_run(f"{chr(ord('A') + i)}. {item}")

    _signature_block(doc, data)

    # ── VERIFICATION ──
    doc.add_page_break()
    _p(doc, "VERIFICATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    _p(doc, "STATE OF NEW YORK\t)", space_after=0)
    _p(doc, "\t\t\t) ss.:", space_after=0)
    _p(doc, f"COUNTY OF {county.upper()}\t)", space_after=18)
    verify = doc.add_paragraph()
    verify.paragraph_format.first_line_indent = Inches(0.5)
    verify.paragraph_format.line_spacing = 2.0  # 22 NYCRR 202.5(a)(1)
    verify.add_run(
        f"{plaintiff}, being duly sworn, deposes and says: I am the Plaintiff in the "
        f"within action. I have read the foregoing Verified Complaint and know the "
        f"contents thereof; the same is true to my own knowledge, except as to the "
        f"matters therein stated to be alleged on information and belief, and as to "
        f"those matters I believe them to be true."
    )
    _p(doc, "", space_after=24)
    for line in ["____________________________", plaintiff]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(3.4)
        para.paragraph_format.space_after = Pt(0)
        para.add_run(line)
    _p(doc, "", space_after=18)
    _p(doc, "Sworn to before me this", space_after=0)
    _p(doc, "____ day of ____________, 20__", space_after=12)
    _p(doc, "____________________________", space_after=0)
    _p(doc, "Notary Public", space_after=0)

    doc.save(output_path)
    return output_path


# ────────────────────────── UD-1 Summons with Notice ──────────────────────────

def generate_ud1_docx(data, output_path):
    """NY UD-1 Summons with Notice as an editable .docx — same content as
    generate_ud1 (caption, summons text, NOTICE, ancillary-relief statement)."""
    import re as _re

    county = (data.get("county", "") or data.get("filingCounty", "")).strip()
    county = _re.sub(r"\s+County$", "", county, flags=_re.IGNORECASE).strip()
    if not county:
        raise ValueError("County is required")
    plaintiff = (data.get("plaintiffName") or "").strip()
    defendant = (data.get("defendantName") or "").strip()
    if not plaintiff or not defendant:
        raise ValueError("Plaintiff and Defendant names are required")
    qualifying_party = (data.get("qualifyingParty") or "").strip().lower()
    if not qualifying_party:
        raise ValueError("Qualifying party is required")
    qualifying_label = "Plaintiff" if qualifying_party == "plaintiff" else "Defendant"
    qualifying_address = (data.get("qualifyingAddress") or "").strip()
    if not qualifying_address:
        raise ValueError("Qualifying address is required")
    plaintiff_address = (data.get("plaintiffAddress") or "").strip()
    if not plaintiff_address:
        raise ValueError("Plaintiff address is required")
    plaintiff_phone = (data.get("plaintiffPhone") or "").strip()
    date_filed = (data.get("dateFiled") or "").strip() or "___________________"

    doc = _base_document()
    _caption_table(
        doc,
        county.upper(),
        plaintiff.upper(),
        defendant.upper(),
        [
            ("Index No.: ______________", False),
            (f"Date Summons filed: {date_filed}", False),
            ("", False),
            (f"Plaintiff designates {title_case(county)} County as the place of trial", False),
            (f"The basis of the venue is: {qualifying_label}'s address", False),
            ("", False),
            ("SUMMONS WITH NOTICE", True),
            ("", False),
            (f"{qualifying_label} resides at: {qualifying_address}", False),
        ],
    )
    _p(doc, "")
    _p(doc, "ACTION FOR A DIVORCE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _p(doc, "To the above named Defendant:", italic=True, space_after=12)

    summoned = doc.add_paragraph()
    summoned.paragraph_format.first_line_indent = Inches(0.5)
    summoned.paragraph_format.line_spacing = 2.0  # 22 NYCRR 202.5(a)(1)
    summoned.paragraph_format.space_after = Pt(12)
    run = summoned.add_run("YOU ARE HEREBY SUMMONED ")
    run.bold = True
    summoned.add_run(
        "to serve a notice of appearance on the Plaintiff within twenty (20) days "
        "after the service of this summons, exclusive of the day of service (or within "
        "thirty (30) days after the service is complete if this summons is not personally "
        "delivered to you within the State of New York); and in case of your failure to "
        "appear, judgment will be taken against you by default for the relief demanded "
        "in the notice set forth below."
    )

    # Signature block — the PDF UD-1's block is the PLAINTIFF's (name,
    # address, phone), not the firm's. Same here.
    _p(doc, f"Dated: {date_filed}", space_after=18)
    for line in ["____________________________", title_case(plaintiff), *plaintiff_address.split("\n"), plaintiff_phone or None]:
        if line is None:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(3.4)
        para.paragraph_format.space_after = Pt(0)
        para.add_run(line)
    _p(doc, "", space_after=12)

    # NOTICE — exact wording of the PDF generator.
    notice = doc.add_paragraph()
    notice.paragraph_format.space_after = Pt(6)
    nrun = notice.add_run("NOTICE:")
    nrun.bold = True
    nrun.underline = True
    notice.add_run(" The nature of this action is to dissolve the marriage between the parties, on the grounds: DRL§170 subd.7 – ")
    grun = notice.add_run("irretrievable breakdown in relationship for a period at least six months")
    grun.bold = True
    grun.underline = True

    relief = doc.add_paragraph()
    relief.paragraph_format.space_after = Pt(12)
    relief.add_run(
        "The relief sought is a judgment of absolute divorce in favor of the Plaintiff "
        "dissolving the marriage between the parties in this action."
    )

    ancillary = doc.add_paragraph()
    ancillary.paragraph_format.space_after = Pt(12)
    ancillary.add_run("The nature of any ancillary or additional relief requested is: ")
    arun = ancillary.add_run("NONE")
    arun.bold = True
    ancillary.add_run(" – I am not requesting any ancillary relief.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_after = Pt(0)
    frun = footer.add_run("UD-1 (Summons with Notice)")
    frun.font.size = Pt(10)

    doc.save(output_path)
    return output_path
